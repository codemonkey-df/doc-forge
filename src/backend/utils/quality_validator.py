"""Quality validator for DOCX output (Story 5.5).

Validates generated DOCX files against FC011 criteria:
- Heading hierarchy (no skipped levels)
- Image rendering (no broken images)
- Code block formatting (monospace fonts)
- Table structure (consistent column counts)
"""

from pathlib import Path
from typing import TypedDict


class QualityResult(TypedDict, total=False):
    """Result of quality validation."""

    passed: bool
    issues: list[str]
    score: int


class QualityValidator:
    """Validates DOCX output against FC011 criteria."""

    ALLOWED_CODE_FONTS = {"Courier New", "Consolas", "Monaco"}

    def validate(self, docx_path: Path) -> QualityResult:
        """Load DOCX and run all FC011 checks.

        Args:
            docx_path: Path to the DOCX file to validate.

        Returns:
            QualityResult with passed, issues, and score.
        """
        # Try to load document
        try:
            from docx import Document
        except ImportError:
            return QualityResult(
                passed=False,
                issues=["Failed to load DOCX: python-docx not installed"],
                score=0,
            )

        if not docx_path.exists():
            return QualityResult(
                passed=False,
                issues=[f"Failed to load DOCX: file not found at {docx_path}"],
                score=0,
            )

        try:
            doc = Document(str(docx_path))
        except Exception as e:
            return QualityResult(
                passed=False,
                issues=[f"Failed to load DOCX: {str(e)}"],
                score=0,
            )

        # Empty document is valid (no spurious issues)
        if not doc.paragraphs and not doc.tables:
            return QualityResult(passed=True, issues=[], score=100)

        # Run all FC011 checks
        issues: list[str] = []

        # Check headings
        heading_issues = self._check_headings(doc)
        issues.extend(heading_issues)

        # Check images
        image_issues = self._check_images(doc)
        issues.extend(image_issues)

        # Check code blocks
        code_issues = self._check_code_blocks(doc)
        issues.extend(code_issues)

        # Check tables
        table_issues = self._check_tables(doc)
        issues.extend(table_issues)

        # Calculate score
        if issues:
            # Deduct 10 points per issue, minimum score is 0
            score = max(0, 100 - (len(issues) * 10))
        else:
            score = 100

        return QualityResult(
            passed=len(issues) == 0,
            issues=issues,
            score=score,
        )

    def _check_headings(self, doc) -> list[str]:
        """Check heading hierarchy for skipped levels.

        Args:
            doc: python-docx Document object.

        Returns:
            List of issues found (empty if valid).
        """
        issues: list[str] = []
        last_level = 0

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Extract heading level from style name (e.g., "Heading 1" -> 1)
                try:
                    level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    continue

                # Check for skipped levels (e.g., Heading 1 -> Heading 3)
                if last_level > 0 and level > last_level + 1:
                    issues.append(
                        f"Skipped heading level: jumped from H{last_level} to H{level}"
                    )

                last_level = level

        return issues

    def _check_images(self, doc) -> list[str]:
        """Check for broken images.

        Args:
            doc: python-docx Document object.

        Returns:
            List of issues found (empty if valid).
        """
        issues: list[str] = []

        # Check for inline images in paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if hasattr(run, "_element") and run._element.xpath(
                    ".//w:drawing/wp:inline"
                ):
                    # Has inline image - check if it has a blip
                    for drawing in run._element.xpath(".//w:drawing"):
                        blips = drawing.xpath(".//a:blip")
                        for blip in blips:
                            embed_attr = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            link_attr = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"
                            )

                            # If neither embed nor link, it's a broken image
                            if not embed_attr and not link_attr:
                                issues.append("Broken image: image reference missing")
                                break

        # Check for images in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if hasattr(run, "_element") and run._element.xpath(
                                ".//w:drawing"
                            ):
                                for drawing in run._element.xpath(".//w:drawing"):
                                    blips = drawing.xpath(".//a:blip")
                                    for blip in blips:
                                        embed_attr = blip.get(
                                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                        )
                                        link_attr = blip.get(
                                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"
                                        )

                                        if not embed_attr and not link_attr:
                                            issues.append(
                                                "Broken image in table: image reference missing"
                                            )
                                            break

        return issues

    def _check_code_blocks(self, doc) -> list[str]:
        """Check code block formatting for monospace fonts.

        Args:
            doc: python-docx Document object.

        Returns:
            List of issues found (empty if valid).
        """
        issues: list[str] = []

        for para in doc.paragraphs:
            # Check if paragraph looks like code (has indentation or is in a code style)
            style_name = para.style.name.lower()

            # Look for code-like paragraphs (indented or code-named styles)
            is_code_like = (
                style_name == "code"
                or style_name.startswith("code")
                or (para.text.startswith("    ") or para.text.startswith("\t"))
            )

            if is_code_like:
                # Check if any run uses a non-monospace font
                for run in para.runs:
                    if run.font.name and run.font.name not in self.ALLOWED_CODE_FONTS:
                        issues.append(
                            f"Code block uses non-monospace font: {run.font.name}"
                        )
                        break

        return issues

    def _check_tables(self, doc) -> list[str]:
        """Check table structure for consistent column counts.

        Args:
            doc: python-docx Document object.

        Returns:
            List of issues found (empty if valid).
        """
        issues: list[str] = []

        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            # Get expected column count from first row
            expected_cols = len(table.rows[0].cells)

            # Check each subsequent row
            for row_idx, row in enumerate(table.rows[1:], start=2):
                actual_cols = len(row.cells)
                if actual_cols != expected_cols:
                    issues.append(
                        f"Inconsistent table columns: table {table_idx + 1}, "
                        f"row {row_idx} has {actual_cols} columns, expected {expected_cols}"
                    )
                    break

        return issues
