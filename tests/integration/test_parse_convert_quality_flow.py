"""Integration tests for parse → convert → quality_check flow (Epics 5.2-5.5).

Tests multi-node flows: parse_to_json → convert_docx → quality_check with mocked subprocess.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from backend.graph_nodes import (
    convert_with_docxjs_node,
    parse_to_json_node,
    quality_check_node,
)
from backend.state import DocumentState


class TestParseConvertQualityFlow:
    """Test parse → convert → quality_check flow with mocked subprocess."""

    @patch("backend.graph_nodes.SessionManager")
    def test_parse_creates_valid_structure_json(
        self,
        mock_sm_class: MagicMock,
        session_with_temp_output: tuple[Path, DocumentState],
    ) -> None:
        """GIVEN valid markdown in temp_output.md WHEN parse_to_json_node runs THEN structure.json written with sections."""
        temp_session_dir, initial_state = session_with_temp_output

        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # Execute parse_to_json_node
        result = parse_to_json_node(initial_state)

        # Assert: structure.json created
        structure_file = temp_session_dir / "structure.json"
        assert structure_file.exists(), "structure.json should be created"

        # Assert: structure has expected keys
        import json

        structure = json.loads(structure_file.read_text(encoding="utf-8"))
        assert "sections" in structure
        assert "metadata" in structure

        # Assert: at least one section (Introduction)
        assert len(structure["sections"]) >= 1

        # Assert: state has structure_json_path set
        assert result["structure_json_path"] == str(structure_file)

    @patch("subprocess.run")
    @patch("backend.graph_nodes.SessionManager")
    def test_parse_then_convert_with_mocked_node(
        self,
        mock_sm_class: MagicMock,
        mock_run: MagicMock,
        session_with_temp_output: tuple[Path, DocumentState],
    ) -> None:
        """GIVEN parse_to_json succeeds WHEN convert_with_docxjs_node runs THEN conversion_success=True."""
        temp_session_dir, initial_state = session_with_temp_output

        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # First parse
        parse_result = parse_to_json_node(initial_state)
        assert parse_result.get("structure_json_path")

        # Mock Node.js converter to succeed
        mock_run.return_value = MagicMock(
            returncode=0, stdout="Conversion complete", stderr=""
        )

        # Execute convert
        result = convert_with_docxjs_node(parse_result)

        # Assert: conversion succeeded
        assert result["conversion_success"] is True
        assert "output_docx_path" in result
        assert result["status"] == "quality_checking"

    @patch("subprocess.run")
    @patch("backend.graph_nodes.SessionManager")
    def test_convert_failure_routes_to_error_handler(
        self,
        mock_sm_class: MagicMock,
        mock_run: MagicMock,
        session_with_temp_output: tuple[Path, DocumentState],
    ) -> None:
        """GIVEN convert fails (mock subprocess rc=1) WHEN convert_with_docxjs_node runs THEN conversion_success=False, routes to error_handler."""
        temp_session_dir, initial_state = session_with_temp_output

        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # First parse
        parse_result = parse_to_json_node(initial_state)

        # Mock Node.js converter to fail
        mock_run.return_value = MagicMock(
            returncode=1, stdout="", stderr="Conversion failed: invalid JSON"
        )

        # Execute convert
        result = convert_with_docxjs_node(parse_result)

        # Assert: conversion failed
        assert result["conversion_success"] is False
        assert result["last_error"] != ""
        assert result["status"] == "error_handling"

    def test_quality_check_with_valid_docx(
        self, temp_session_dir: Path, sample_state: DocumentState
    ) -> None:
        """GIVEN valid DOCX file WHEN quality_check_node runs THEN quality_passed=True."""
        # Create a simple valid DOCX using python-docx
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create DOCX
        docx_path = temp_session_dir / "output.docx"
        doc = DocxDocument()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test document.")
        doc.save(str(docx_path))

        # Execute quality_check
        state = sample_state.copy()
        state["output_docx_path"] = str(docx_path)
        result = quality_check_node(state)

        # Assert: quality passed
        assert result["quality_passed"] is True
        assert result["status"] == "complete"
        assert len(result.get("quality_issues", [])) == 0

    def test_quality_check_with_missing_docx(
        self, temp_session_dir: Path, sample_state: DocumentState
    ) -> None:
        """GIVEN missing DOCX file WHEN quality_check_node runs THEN quality_passed=False."""
        # Execute quality_check with no DOCX
        state = sample_state.copy()
        state["output_docx_path"] = ""
        result = quality_check_node(state)

        # Assert: quality failed
        assert result["quality_passed"] is False
        assert "No DOCX output to validate" in result["last_error"]
        assert result["status"] == "error_handling"


class TestParseErrorHandling:
    """Test parse_to_json_node error handling."""

    @patch("backend.graph_nodes.SessionManager")
    def test_parse_with_missing_temp_md_creates_empty_structure(
        self,
        mock_sm_class: MagicMock,
        temp_session_dir: Path,
        sample_state: DocumentState,
    ) -> None:
        """GIVEN temp_output.md missing WHEN parse_to_json_node runs THEN creates empty structure.json."""
        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # Don't create temp_output.md - leave it missing

        # Execute parse
        result = parse_to_json_node(sample_state)

        # Assert: structure.json created with empty sections
        structure_file = temp_session_dir / "structure.json"
        assert structure_file.exists()

        import json

        structure = json.loads(structure_file.read_text(encoding="utf-8"))
        assert structure["sections"] == []

        # Assert: state has structure_json_path
        assert result["structure_json_path"] == str(structure_file)


class TestConvertRouting:
    """Test routing decisions after conversion."""

    @patch("subprocess.run")
    @patch("backend.graph_nodes.SessionManager")
    def test_convert_success_routes_to_quality_check(
        self,
        mock_sm_class: MagicMock,
        mock_run: MagicMock,
        session_with_temp_output: tuple[Path, DocumentState],
    ) -> None:
        """GIVEN conversion_success=True WHEN route_after_conversion THEN routes to quality_check."""
        temp_session_dir, initial_state = session_with_temp_output

        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # Parse first
        parse_result = parse_to_json_node(initial_state)

        # Mock Node.js to succeed
        mock_run.return_value = MagicMock(returncode=0, stdout="ok", stderr="")

        # Convert
        convert_result = convert_with_docxjs_node(parse_result)

        # Assert: would route to quality_check
        assert convert_result["conversion_success"] is True
        assert convert_result["status"] == "quality_checking"

    @patch("subprocess.run")
    @patch("backend.graph_nodes.SessionManager")
    def test_convert_failure_routes_to_error_handler(
        self,
        mock_sm_class: MagicMock,
        mock_run: MagicMock,
        session_with_temp_output: tuple[Path, DocumentState],
    ) -> None:
        """GIVEN conversion_success=False WHEN route_after_conversion THEN routes to error_handler."""
        temp_session_dir, initial_state = session_with_temp_output

        # Mock SessionManager
        mock_sm = MagicMock()
        mock_sm.get_path.return_value = temp_session_dir
        mock_sm_class.return_value = mock_sm

        # Parse first
        parse_result = parse_to_json_node(initial_state)

        # Mock Node.js to fail - but set conversion_success=True first
        # This test just verifies the status when conversion fails
        mock_run.return_value = MagicMock(returncode=1, stdout="", stderr="Error")

        # Convert
        convert_result = convert_with_docxjs_node(parse_result)

        # Assert: conversion failed
        assert convert_result["conversion_success"] is False
        assert convert_result["status"] == "error_handling"
